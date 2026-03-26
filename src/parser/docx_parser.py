"""DOCX / DOC parser: handles paragraphs, tables, headers, footers, and textboxes.

Supported formats
-----------------
- .docx  — handled directly by python-docx.
- .doc   — converted to a temporary .docx first via macOS ``textutil``
             (built-in on every Mac, no extra install needed) or
             LibreOffice if available.  The translated file is always
             saved as .docx regardless of input extension.

Design decisions
----------------
- Each non-empty paragraph → one ContentBlock (merge runs into plain text).
- On rebuild: write translated text into the first run, clear remaining runs
  (preserves the first run's character format: bold, italic, font size, color).
- Skip paragraphs that are clearly non-translatable (pure numbers, URLs, code).
- Text inside text boxes (``w:txbxContent``) is extracted and translated via
  lxml XPath since python-docx does not expose textbox content.
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from lxml import etree

from src.models.content import BlockType, ContentBlock, FileMeta, ParsedFile
from src.parser.base import BaseParser

logger = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_NSMAP = {"w": _W_NS}

# Extended namespace map for textbox discovery via XPath.
_TXBX_NSMAP = {
    "w": _W_NS,
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "v": "urn:schemas-microsoft-com:vml",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}

# XPath to find all w:txbxContent elements (modern *and* legacy VML paths).
_TXBX_CONTENT_XPATH = (
    ".//wps:txbx/w:txbxContent | .//v:textbox/w:txbxContent"
)


def _xpath_txbx(body_elem) -> list:
    """Find all ``w:txbxContent`` elements via lxml xpath.

    python-docx's ``BaseOxmlElement.xpath()`` only supports the built-in
    namespace map which lacks ``wps`` / ``v``.  We call lxml's raw xpath
    via ``etree.ElementBase`` to use a custom namespace map.
    """
    return etree.ElementBase.xpath(
        body_elem, _TXBX_CONTENT_XPATH, namespaces=_TXBX_NSMAP
    )


def _run_has_drawing(run) -> bool:
    """Return True if the run contains a ``w:drawing`` element (image)."""
    return len(run._r.findall("w:drawing", _W_NSMAP)) > 0


def _safe_set_run_text(run, text: str) -> None:
    """Set the text of a run using direct lxml operations.

    Unlike ``run.text = ...`` this only removes ``w:t``, ``w:br`` and
    ``w:tab`` children — ``w:drawing`` (images) and ``w:rPr`` (formatting)
    are left intact.
    """
    r_elem = run._r
    for tag in ("w:t", "w:br", "w:tab"):
        for child in r_elem.findall(tag, _W_NSMAP):
            r_elem.remove(child)
    if text:
        t = etree.SubElement(r_elem, f"{{{_W_NS}}}t")
        t.text = text
        if text.startswith(" ") or text.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _safe_clear_run_text(run) -> None:
    """Remove text from a run while preserving images and formatting."""
    _safe_set_run_text(run, "")

_URL_RE = re.compile(r"^https?://\S+$|^www\.\S+$", re.IGNORECASE)
_NUM_ONLY_RE = re.compile(r"^[\d\s.,;:+\-*/=%$€£¥()[\]{}]+$")
_BULLET_PREFIX_RE = re.compile(r"^([•●▪■◆◇○◦·\-*]\s+)(.+)$")
_NUMBER_PREFIX_RE = re.compile(
    r"^((?:\(?\d+[\).、]|[A-Za-z][\).、]|[IVXLCMivxlcm]+[\).、])\s+)(.+)$"
)
_LIST_STYLE_KEYWORDS = ("list", "bullet", "number", "编号", "项目符号")


def _is_translatable(text: str) -> bool:
    """Return False for content that should not be translated."""
    t = text.strip()
    if not t:
        return False
    if len(t) < 2:
        return False
    if _URL_RE.match(t):
        return False
    if _NUM_ONLY_RE.match(t):
        return False
    return True


def _convert_doc_to_docx(doc_path: str) -> str:
    """Convert a legacy .doc file to .docx and return the new path.

    Tries, in order:
    1. macOS ``textutil`` (built-in on every Mac — no extra install).
    2. ``libreoffice --headless --convert-to docx`` (if LibreOffice is installed).
    3. ``soffice`` (alternate LibreOffice entry point).

    Raises ``RuntimeError`` if no converter is available.
    """
    source_path = Path(doc_path)
    out_dir = Path(
        tempfile.mkdtemp(
            prefix="agent_translation_doc_",
            dir=str(source_path.parent),
        )
    )
    out_docx = out_dir / f"{source_path.stem}.docx"

    # ── 1 / 2. LibreOffice (better list fidelity than textutil) ──────
    for cmd in ("libreoffice", "soffice"):
        try:
            result = subprocess.run(
                [cmd, "--headless", "--convert-to", "docx",
                 "--outdir", str(out_dir), doc_path],
                capture_output=True,
                timeout=120,
            )
            if result.returncode == 0 and out_docx.exists():
                logger.info(
                    "Converted %s → %s via %s", doc_path, out_docx, cmd
                )
                return str(out_docx)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # ── 3. macOS textutil ─────────────────────────────────────────────
    try:
        result = subprocess.run(
            ["textutil", "-convert", "docx", "-output", str(out_docx), doc_path],
            capture_output=True,
            timeout=60,
        )
        if result.returncode == 0 and out_docx.exists():
            logger.info("Converted %s → %s via textutil", doc_path, out_docx)
            return str(out_docx)
        logger.debug(
            "textutil failed (rc=%d): %s",
            result.returncode,
            result.stderr.decode(errors="replace"),
        )
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.debug("textutil not available: %s", e)

    shutil.rmtree(out_dir, ignore_errors=True)
    raise RuntimeError(
        f"无法转换 .doc 文件：{Path(doc_path).name}\n"
        "请尝试以下任一方案：\n"
        "  ① 用 Word / WPS 将文件另存为 .docx 格式后再上传；\n"
        "  ② 安装 LibreOffice（https://www.libreoffice.org/）。"
    )


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents (.docx and legacy .doc)."""

    EXTENSIONS = {".docx", ".doc"}

    def can_handle(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.EXTENSIONS

    # ------------------------------------------------------------------
    # parse
    # ------------------------------------------------------------------
    def parse(self, file_path: str) -> ParsedFile:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX support. "
                "Run: pip install python-docx>=1.1"
            ) from exc

        suffix = Path(file_path).suffix.lower()

        if suffix == ".doc":
            converted = _convert_doc_to_docx(file_path)
            actual_path = converted
        else:
            actual_path = file_path

        doc = Document(actual_path)
        blocks: list[ContentBlock] = []
        word_count = 0

        # ── body paragraphs ──────────────────────────────────────────
        for i, para in enumerate(doc.paragraphs):
            block = self._para_to_block(para, f"p_{i}")
            if block:
                blocks.append(block)
                word_count += len(block.source_text.split())

        # ── tables ───────────────────────────────────────────────────
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        bid = f"tbl{t_idx}_r{r_idx}_c{c_idx}_p{p_idx}"
                        block = self._para_to_block(
                            para,
                            bid,
                            extra_meta={
                                "in_table": True,
                                "table_idx": t_idx,
                                "row": r_idx,
                                "col": c_idx,
                            },
                        )
                        if block:
                            blocks.append(block)
                            word_count += len(block.source_text.split())

        # ── headers / footers ────────────────────────────────────────
        for sec_idx, section in enumerate(doc.sections):
            for hf_name, hf_obj in (
                ("header", section.header),
                ("footer", section.footer),
            ):
                if hf_obj is None:
                    continue
                try:
                    for p_idx, para in enumerate(hf_obj.paragraphs):
                        bid = f"sec{sec_idx}_{hf_name}_p{p_idx}"
                        block = self._para_to_block(
                            para, bid, extra_meta={"section": hf_name}
                        )
                        if block:
                            blocks.append(block)
                except Exception:
                    logger.debug("Could not parse %s in section %d", hf_name, sec_idx)

        # ── text boxes ────────────────────────────────────────────────
        # python-docx does not expose textbox content; use XPath to find
        # all w:txbxContent elements (modern wps:txbx and legacy v:textbox).
        try:
            from docx.text.paragraph import Paragraph

            for tb_idx, txbx_content in enumerate(_xpath_txbx(doc.element.body)):
                p_elems = txbx_content.findall(f"{{{_W_NS}}}p")
                for p_idx, p_elem in enumerate(p_elems):
                    para = Paragraph(p_elem, doc.part)
                    bid = f"txbx{tb_idx}_p{p_idx}"
                    block = self._para_to_block(
                        para, bid, extra_meta={"in_textbox": True}
                    )
                    if block:
                        blocks.append(block)
                        word_count += len(block.source_text.split())
        except Exception:
            logger.debug("Could not parse textbox content", exc_info=True)

        # For .doc input, record the converted file as template so rebuild
        # can use it; we also store the source suffix for reference.
        return ParsedFile(
            meta=FileMeta(
                original_name=os.path.basename(file_path),
                file_type=suffix.lstrip("."),
                word_count=word_count,
            ),
            blocks=blocks,
            format_template=str(Path(actual_path).resolve()),
        )

    @staticmethod
    def _para_to_block(
        para, block_id: str, extra_meta: dict | None = None
    ) -> ContentBlock | None:
        text = para.text.strip()
        style_name = (para.style.name if para.style else "") or ""
        list_meta = DocxParser._extract_list_meta(para, text, style_name)
        normalized_text = list_meta["content_text"]

        if not _is_translatable(normalized_text):
            return None

        # Detect block type from paragraph style name / list metadata
        style_name_lower = style_name.lower()
        if list_meta["is_list"]:
            block_type = BlockType.LIST
        elif "heading" in style_name_lower or "title" in style_name_lower:
            block_type = BlockType.HEADING
        else:
            block_type = BlockType.PARAGRAPH

        # Save first run's format for rebuild
        first_run_fmt: dict = {}
        if para.runs:
            r = para.runs[0]
            first_run_fmt = {
                "bold": r.bold,
                "italic": r.italic,
                "font_name": r.font.name,
                "font_size": str(r.font.size) if r.font.size else None,
            }

        meta: dict = {
            "style_name": style_name,
            "first_run_fmt": first_run_fmt,
            **list_meta,
        }
        if extra_meta:
            meta.update(extra_meta)

        return ContentBlock(
            id=block_id,
            type=block_type,
            source_text=normalized_text,
            metadata=meta,
        )

    @staticmethod
    def _extract_list_meta(para, text: str, style_name: str) -> dict:
        prefix = ""
        content_text = text
        style_name_lower = style_name.lower()
        list_style = any(keyword in style_name_lower for keyword in _LIST_STYLE_KEYWORDS)

        number_match = _NUMBER_PREFIX_RE.match(text)
        bullet_match = _BULLET_PREFIX_RE.match(text)
        if number_match:
            prefix = number_match.group(1)
            content_text = number_match.group(2).strip()
        elif bullet_match:
            prefix = bullet_match.group(1)
            content_text = bullet_match.group(2).strip()

        list_level = None
        list_num_id = None
        p_pr = getattr(getattr(para, "_p", None), "pPr", None)
        num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None
        if num_pr is not None:
            ilvl = getattr(num_pr, "ilvl", None)
            num_id = getattr(num_pr, "numId", None)
            if ilvl is not None and getattr(ilvl, "val", None) is not None:
                list_level = int(ilvl.val)
            if num_id is not None and getattr(num_id, "val", None) is not None:
                list_num_id = int(num_id.val)

        is_list = bool(prefix or list_style or list_num_id is not None)
        return {
            "is_list": is_list,
            "list_prefix": prefix,
            "list_level": list_level,
            "list_num_id": list_num_id,
            "content_text": content_text,
        }

    # ------------------------------------------------------------------
    # rebuild
    # ------------------------------------------------------------------
    def rebuild(self, parsed_file: ParsedFile, output_path: str) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("python-docx is required for DOCX support.") from exc

        original_path = parsed_file.format_template
        if not original_path or not os.path.exists(original_path):
            raise FileNotFoundError(f"Original DOCX not found: {original_path}")

        # .doc inputs were converted to .docx for parsing; the template already
        # points to that .docx.  Always output as .docx regardless of source ext.
        if Path(output_path).suffix.lower() == ".doc":
            output_path = str(Path(output_path).with_suffix(".docx"))

        doc = Document(original_path)
        block_map: dict[str, ContentBlock] = {b.id: b for b in parsed_file.blocks}

        # ── body paragraphs ──────────────────────────────────────────
        for i, para in enumerate(doc.paragraphs):
            block = block_map.get(f"p_{i}")
            if block:
                self._write_back(para, self._compose_output_text(block))

        # ── tables ───────────────────────────────────────────────────
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        bid = f"tbl{t_idx}_r{r_idx}_c{c_idx}_p{p_idx}"
                        block = block_map.get(bid)
                        if block:
                            self._write_back(para, self._compose_output_text(block))

        # ── headers / footers ────────────────────────────────────────
        for sec_idx, section in enumerate(doc.sections):
            for hf_name, hf_obj in (
                ("header", section.header),
                ("footer", section.footer),
            ):
                if hf_obj is None:
                    continue
                try:
                    for p_idx, para in enumerate(hf_obj.paragraphs):
                        bid = f"sec{sec_idx}_{hf_name}_p{p_idx}"
                        block = block_map.get(bid)
                        if block:
                            self._write_back(para, self._compose_output_text(block))
                except Exception:
                    pass

        # ── text boxes ────────────────────────────────────────────────
        try:
            from docx.text.paragraph import Paragraph

            for tb_idx, txbx_content in enumerate(_xpath_txbx(doc.element.body)):
                p_elems = txbx_content.findall(f"{{{_W_NS}}}p")
                for p_idx, p_elem in enumerate(p_elems):
                    bid = f"txbx{tb_idx}_p{p_idx}"
                    block = block_map.get(bid)
                    if block:
                        para = Paragraph(p_elem, doc.part)
                        self._write_back(para, self._compose_output_text(block))
        except Exception:
            logger.debug("Could not rebuild textbox content", exc_info=True)

        doc.save(output_path)
        logger.info("Rebuilt DOCX saved to %s", output_path)
        return output_path

    @staticmethod
    def _compose_output_text(block: ContentBlock) -> str:
        text = BaseParser._best_text(block)
        prefix = str(block.metadata.get("list_prefix") or "")
        if prefix and text and not text.startswith(prefix):
            return f"{prefix}{text}"
        return text

    @staticmethod
    def _write_back(para, translated_text: str) -> None:
        """Write translated text into the paragraph.

        Strategy: write full translated text into the first *text-bearing*
        run using safe lxml operations, then clear text from remaining
        runs.  Runs that contain ``w:drawing`` elements (images) are
        preserved — only their text is stripped so the image stays in
        place.
        """
        if not para.runs:
            # No runs — check for images before clearing
            p_elem = para._p
            has_drawing = len(
                p_elem.findall(".//w:drawing", _W_NSMAP)
            ) > 0
            if has_drawing:
                # Keep the images, just append a new run with the text
                para.add_run(translated_text)
            else:
                try:
                    if para.text.strip():
                        para.clear()
                        para.add_run(translated_text)
                except Exception:
                    pass
            return

        # Write translated text into the first run (safe: keeps images)
        _safe_set_run_text(para.runs[0], translated_text)

        # Clear text from subsequent runs, but keep image-bearing runs
        for run in para.runs[1:]:
            if _run_has_drawing(run):
                _safe_clear_run_text(run)
            else:
                _safe_set_run_text(run, "")
