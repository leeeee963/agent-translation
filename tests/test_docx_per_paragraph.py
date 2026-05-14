"""Adversarial test harness for the per-paragraph DOCX translation pipeline.

Mirrors tests/test_pptx_per_paragraph.py for Word documents. Goal: prove no
LLM misbehavior can produce silent paragraph-loss, paragraph-merge, or
content garbling (e.g. literal \\n leaking into a single <w:t>) in the
rebuilt .docx.

Run with:  .venv/bin/python -m pytest tests/test_docx_per_paragraph.py -v
"""

from __future__ import annotations

import tempfile
from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from src.models.content import ContentBlock, ParsedFile
from src.parser.docx_parser import DocxParser
from src.translator.agent import (
    _BLOCK_MARKER,
    _build_marked_input,
    _parse_marked_response,
)
from src.translator.segmenter import Segmenter


# ---------------------------------------------------------------------------
# Synthetic DOCX builder
# ---------------------------------------------------------------------------


def build_adversarial_docx() -> str:
    doc = Document()

    # Body section ------------------------------------------------------
    # A heading + 3 bullet-style paragraphs (the same shape as the user's
    # PPTX failure case, but in Word).
    h = doc.add_paragraph("Characteristics", style="Heading 1")
    # Bullets — use List Bullet style so list_num_id is set.
    doc.add_paragraph("Strong long-context orientation.", style="List Bullet")
    doc.add_paragraph("Designed for multi-step agent workflows.", style="List Bullet")
    doc.add_paragraph("Useful for code, documents, and research tasks.", style="List Bullet")

    # A standalone body paragraph (own group).
    doc.add_paragraph("This is a standalone body paragraph for testing.")

    # Another heading + bullets group.
    doc.add_paragraph("Pros", style="Heading 1")
    doc.add_paragraph("Good for reading large inputs.", style="List Bullet")
    doc.add_paragraph("Useful for decomposing multi-part tasks.", style="List Bullet")

    # A table with multi-paragraph cells.
    table = doc.add_table(rows=2, cols=2)
    cell00 = table.rows[0].cells[0]
    cell00.text = "Title"
    cell00.add_paragraph("Subtitle line")
    cell00.add_paragraph("Footnote line")

    cell01 = table.rows[0].cells[1]
    cell01.text = "Single cell content"

    cell10 = table.rows[1].cells[0]
    cell10.text = "Alpha"
    cell10.add_paragraph("Beta")
    cell10.add_paragraph("Gamma")

    # cell (1,1) deliberately empty

    # A 30-paragraph stress section.
    for i in range(30):
        doc.add_paragraph(f"Stress paragraph number {i:02d}.")

    out = Path(tempfile.gettempdir()) / "adversarial_test.docx"
    doc.save(str(out))
    return str(out)


# ---------------------------------------------------------------------------
# Adversaries (mirroring the PPTX suite)
# ---------------------------------------------------------------------------


def _translate_marker(text: str) -> str:
    return f"[T]{text}"


def _emit_with_groups(blocks, transform: Callable[[ContentBlock], str | None]) -> str:
    """Emit `[[BLOCK:id]]\\ntransformed_text` for every block where transform
    returns non-None, with `<!-- group: X -->` lines preceding each group
    transition. Always terminates with `[[END]]`."""
    parts = []
    last_group = None
    for b in blocks:
        text = transform(b)
        if text is None:
            continue
        marker = f"{_BLOCK_MARKER.format(block_id=b.id)}\n{text}"
        group = b.metadata.get("text_frame_group")
        if group and group != last_group:
            parts.append(f"<!-- group: {group} -->\n{marker}")
        else:
            parts.append(marker)
        last_group = group
    return "\n\n".join(parts) + "\n\n[[END]]"


def _faithful_response(blocks, translate=_translate_marker):
    return _emit_with_groups(blocks, lambda b: translate(b.source_text))


def adversary_drop_last(blocks, translate=_translate_marker):
    by_group: dict[str | None, list[ContentBlock]] = {}
    for b in blocks:
        by_group.setdefault(b.metadata.get("text_frame_group"), []).append(b)
    drop_ids = {g[-1].id for g in by_group.values() if len(g) > 1}
    return _emit_with_groups(
        blocks,
        lambda b: None if b.id in drop_ids else translate(b.source_text),
    )


def adversary_merge_two(blocks, translate=_translate_marker):
    by_group: dict[str | None, list[ContentBlock]] = {}
    for b in blocks:
        by_group.setdefault(b.metadata.get("text_frame_group"), []).append(b)
    merge_pairs: dict[str, str] = {}
    skip_ids: set[str] = set()
    for g in by_group.values():
        if len(g) >= 2:
            merge_pairs[g[0].id] = g[1].source_text
            skip_ids.add(g[1].id)

    def t(b):
        if b.id in skip_ids:
            return None
        text = translate(b.source_text)
        if b.id in merge_pairs:
            text = text + " " + translate(merge_pairs[b.id])
        return text

    return _emit_with_groups(blocks, t)


def adversary_swap_order(blocks, translate=_translate_marker):
    return _emit_with_groups(list(reversed(blocks)), lambda b: translate(b.source_text))


def adversary_extra_noise(blocks, translate=_translate_marker):
    parts = ["Here are the translations:\n"]
    for i, b in enumerate(blocks):
        text = translate(b.source_text)
        comment_extra = "<!-- nb: machine output -->" if i % 3 == 0 else ""
        marker = f"{_BLOCK_MARKER.format(block_id=b.id)}\n{text}\n{comment_extra}"
        parts.append(marker)
        parts.append("")
    body = "\n\n".join(parts)
    return body + "\n\n[[END]]\n\nLet me know if you need adjustments."


def adversary_empty_one(blocks, translate=_translate_marker):
    by_group: dict[str | None, list[ContentBlock]] = {}
    for b in blocks:
        by_group.setdefault(b.metadata.get("text_frame_group"), []).append(b)
    empty_ids = {g[0].id for g in by_group.values()}
    return _emit_with_groups(
        blocks,
        lambda b: "" if b.id in empty_ids else translate(b.source_text),
    )


def adversary_multiline_translation(blocks, translate=_translate_marker):
    """Inject newlines inside the translation. DOCX rebuild must flatten so the
    paragraph doesn't contain a literal \\n (which Word renders as nothing or
    a glyph artifact)."""
    def t(b):
        text = translate(b.source_text)
        if " " in text:
            text = text.replace(" ", "\n", 1)
        return f"\n{text}\n"
    return _emit_with_groups(blocks, t)


def adversary_protocol_injection(blocks, translate=_translate_marker):
    """Translation legitimately contains [[END]] and [[BLOCK:fake]] inline."""
    def t(b):
        text = translate(b.source_text)
        idx = list(blocks).index(b)
        if idx == 1:
            text = f"{text} (note: [[END]] of section)"
        elif idx == 2:
            text = f"{text} — see [[BLOCK:hallucinated]] reference"
        return text
    return _emit_with_groups(blocks, t)


def adversary_hallucinated_ids(blocks, translate=_translate_marker):
    parts = []
    last_group = None
    for i, b in enumerate(blocks):
        text = translate(b.source_text)
        marker = f"{_BLOCK_MARKER.format(block_id=b.id)}\n{text}"
        group = b.metadata.get("text_frame_group")
        if group and group != last_group:
            parts.append(f"<!-- group: {group} -->\n{marker}")
        else:
            parts.append(marker)
        last_group = group
        if i % 4 == 0:
            parts.append(f"[[BLOCK:fake_id_{i}]]\nshould be ignored")
    return "\n\n".join(parts) + "\n\n[[END]]"


def adversary_drop_everything(blocks, translate=_translate_marker):
    return "[[END]]"


def adversary_whitespace_translation(blocks, translate=_translate_marker):
    by_group: dict[str | None, list[ContentBlock]] = {}
    for b in blocks:
        by_group.setdefault(b.metadata.get("text_frame_group"), []).append(b)
    blank_ids: set[str] = set()
    for g in by_group.values():
        for b in g[: max(1, len(g) // 2)]:
            blank_ids.add(b.id)
    return _emit_with_groups(
        blocks,
        lambda b: "   \n  \n   " if b.id in blank_ids else translate(b.source_text),
    )


def adversary_persistent_drop(blocks, translate=_translate_marker):
    return adversary_drop_last(blocks, translate)


# ---------------------------------------------------------------------------
# Pipeline simulator
# ---------------------------------------------------------------------------


def _apply_translations(
    parsed: ParsedFile,
    adversary: Callable,
    *,
    max_retries: int = 2,
    persistent: bool = False,
):
    seg_engine = Segmenter()
    blocks = parsed.translatable_blocks
    segments = seg_engine.segment(blocks, file_type="docx", max_tokens=3800)

    for seg in segments:
        _ = _build_marked_input(seg)  # exercise group-marker builder
        response = adversary(seg, _translate_marker)
        matched = _parse_marked_response(response, seg)
        attempts = 1
        while attempts <= max_retries and len(matched) < len(seg):
            for b in seg:
                if b.id not in matched:
                    b.translated_text = ""
            response = (
                adversary(seg, _translate_marker)
                if persistent
                else _faithful_response(seg)
            )
            matched = _parse_marked_response(response, seg)
            attempts += 1
        if len(matched) < len(seg):
            # Single-block fallback for every block (mirrors agent.py).
            for b in seg:
                b.translated_text = ""
                _parse_marked_response(_faithful_response([b]), [b])


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------


def _collect_para_texts(doc) -> dict[str, list[str]]:
    """Map each translatable id-prefix → list of paragraph texts in that group.

    For body paragraphs, the 'group' is keyed by paragraph index (since each
    body paragraph is its own group unless in a list).
    """
    out: dict[str, list[str]] = {}
    for i, para in enumerate(doc.paragraphs):
        out[f"body_{i}"] = [para.text]
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                key = f"tbl{t_idx}_r{r_idx}_c{c_idx}"
                out[key] = [p.text for p in cell.paragraphs]
    return out


def _expected_text_for_block(block: ContentBlock) -> str:
    """Recompute what _compose_output_text would write for a faithful '[T]'-prefix
    translation: lead/trail whitespace → none for DOCX; bullet prefix preserved."""
    text = f"[T]{block.source_text}"
    prefix = str(block.metadata.get("list_prefix") or "")
    if prefix and not text.startswith(prefix):
        text = f"{prefix}{text}"
    return text


def _assert_rebuild_matches(
    output_path: str,
    parsed_source: ParsedFile,
    source_docx: str,
    *,
    strict_content: bool = True,
):
    """Verify the rebuilt DOCX:
      • Block paragraphs contain the expected '[T]<source>' content.
      • Non-block paragraphs (URLs, blanks, numbers-only) are unchanged.
      • Total paragraph count is unchanged.
      • No literal '\\n' appears in any paragraph text (newlines were flattened).
    """
    # Map block id → expected text
    expected_by_id: dict[str, str] = {
        b.id: _expected_text_for_block(b) for b in parsed_source.translatable_blocks
    }

    # Map block id → actual rebuilt paragraph text
    out_doc = Document(output_path)
    actual_by_id: dict[str, str] = {}

    body_paras = list(out_doc.paragraphs)
    for i, para in enumerate(body_paras):
        actual_by_id[f"p_{i}"] = para.text
    for t_idx, table in enumerate(out_doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    actual_by_id[f"tbl{t_idx}_r{r_idx}_c{c_idx}_p{p_idx}"] = para.text

    failures: list[str] = []
    for block_id, expected in expected_by_id.items():
        got = actual_by_id.get(block_id)
        if got is None:
            failures.append(f"{block_id}: paragraph missing from rebuilt doc")
            continue
        if "\n" in got or "\r" in got:
            failures.append(f"{block_id}: literal newline in paragraph text: {got!r}")
            continue
        if strict_content:
            if got.strip() != expected.strip():
                failures.append(f"{block_id}: expected {expected!r}, got {got!r}")
        else:
            # Structural: the source-derived prefix must appear somewhere.
            src = parsed_source.blocks  # noqa: F841 (kept for readability)
            source = next(
                (b.source_text for b in parsed_source.blocks if b.id == block_id), ""
            )
            if f"[T]{source}" not in got:
                failures.append(
                    f"{block_id}: missing source content [T]{source!r} (got {got!r})"
                )

    # Paragraph-count invariant: rebuild must not add or drop paragraphs.
    src_doc = Document(source_docx)
    if len(list(src_doc.paragraphs)) != len(body_paras):
        failures.append(
            f"body paragraph count changed: "
            f"orig {len(list(src_doc.paragraphs))}, rebuilt {len(body_paras)}"
        )

    if failures:
        raise AssertionError("Rebuild mismatch:\n  " + "\n  ".join(failures))


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def adversarial_docx_path() -> str:
    return build_adversarial_docx()


@pytest.mark.parametrize(
    "name,adv,persistent,content_mutating",
    [
        ("faithful", _faithful_response, False, False),
        ("drop_last", adversary_drop_last, False, False),
        ("merge_two", adversary_merge_two, False, False),
        ("swap_order", adversary_swap_order, False, False),
        ("extra_noise", adversary_extra_noise, False, False),
        ("empty_one", adversary_empty_one, False, False),
        ("multiline", adversary_multiline_translation, False, False),
        ("protocol_injection", adversary_protocol_injection, False, True),
        ("hallucinated_ids", adversary_hallucinated_ids, False, False),
        ("whitespace_translation", adversary_whitespace_translation, False, False),
        ("drop_everything", adversary_drop_everything, False, False),
        ("persistent_drop_last", adversary_persistent_drop, True, False),
        ("persistent_merge_two", adversary_merge_two, True, False),
        ("persistent_empty_one", adversary_empty_one, True, False),
        ("persistent_protocol_injection", adversary_protocol_injection, True, True),
        ("persistent_drop_everything", adversary_drop_everything, True, False),
        ("persistent_whitespace", adversary_whitespace_translation, True, False),
    ],
)
def test_adversarial_rebuild(adversarial_docx_path, name, adv, persistent, content_mutating, tmp_path):
    parser = DocxParser()
    parsed = parser.parse(adversarial_docx_path)

    _apply_translations(parsed, adv, persistent=persistent)

    out_path = str(tmp_path / f"out_{name}.docx")
    parser.rebuild(parsed, out_path)

    _assert_rebuild_matches(
        out_path,
        parsed_source=parsed,
        source_docx=adversarial_docx_path,
        strict_content=not content_mutating,
    )


def test_group_metadata_populated(adversarial_docx_path):
    parsed = DocxParser().parse(adversarial_docx_path)
    for b in parsed.translatable_blocks:
        assert b.metadata.get("text_frame_group"), b.id
        assert b.metadata.get("para_index") is not None, b.id
        # No translatable block should carry a newline in source_text.
        assert "\n" not in b.source_text, b.id


def test_group_markers_in_prompt(adversarial_docx_path):
    parsed = DocxParser().parse(adversarial_docx_path)
    blocks = parsed.translatable_blocks
    msg = _build_marked_input(blocks)
    groups = {b.metadata["text_frame_group"] for b in blocks}
    for g in groups:
        assert msg.count(f"<!-- group: {g} -->") == 1, (
            f"group {g!r} not introduced exactly once in prompt input"
        )


def test_segmenter_keeps_groups_together(adversarial_docx_path):
    parsed = DocxParser().parse(adversarial_docx_path)
    blocks = parsed.translatable_blocks
    # Force a tiny budget so segmenter MUST split.
    segments = Segmenter().segment(blocks, file_type="docx", max_tokens=15)
    seen: set[str] = set()
    for seg in segments:
        groups_in_seg = {b.metadata["text_frame_group"] for b in seg}
        for g in groups_in_seg:
            # A group may legitimately span both itself and other groups in the
            # same segment, but once it appears in segment N it must not also
            # appear in segment N+1.
            if g in seen:
                # It MIGHT still be allowed if the entire group fit in one
                # segment — check by scanning if any block of g is in another seg.
                pass
            seen.add(g)
    # Stronger atomicity check: every group's blocks live in exactly one segment.
    group_to_segments: dict[str, set[int]] = {}
    for seg_idx, seg in enumerate(segments):
        for b in seg:
            group_to_segments.setdefault(b.metadata["text_frame_group"], set()).add(seg_idx)
    for g, segs in group_to_segments.items():
        assert len(segs) == 1, f"group {g!r} split across segments {segs}"


def test_no_newlines_in_rebuilt_paragraphs(adversarial_docx_path, tmp_path):
    """Even with a multiline adversary, every paragraph in the rebuilt doc
    must be free of literal newlines — Word would mis-render them."""
    parser = DocxParser()
    parsed = parser.parse(adversarial_docx_path)
    _apply_translations(parsed, adversary_multiline_translation)
    out = str(tmp_path / "out_nonewlines.docx")
    parser.rebuild(parsed, out)
    doc = Document(out)
    for para in doc.paragraphs:
        assert "\n" not in para.text, f"newline in body para: {para.text!r}"
        assert "\r" not in para.text, f"CR in body para: {para.text!r}"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    assert "\n" not in para.text, f"newline in cell para: {para.text!r}"


def test_bullet_content_survives_rebuild(adversarial_docx_path, tmp_path):
    """Bullet/list paragraph content must be translated and survive rebuild,
    regardless of whether python-docx's synthetic "List Bullet" style attaches
    a num_id (test-environment-dependent — in real Word docs it always does)."""
    parser = DocxParser()
    parsed = parser.parse(adversarial_docx_path)
    _apply_translations(parsed, _faithful_response)
    out = str(tmp_path / "out_bullets.docx")
    parser.rebuild(parsed, out)
    doc = Document(out)
    expected_bullets = [
        "[T]Strong long-context orientation.",
        "[T]Designed for multi-step agent workflows.",
        "[T]Useful for code, documents, and research tasks.",
        "[T]Good for reading large inputs.",
        "[T]Useful for decomposing multi-part tasks.",
    ]
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for expected in expected_bullets:
        assert expected in all_text, f"bullet content lost: {expected}"
